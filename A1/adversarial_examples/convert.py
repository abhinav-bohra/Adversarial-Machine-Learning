import docx
import pandas as pd
from docx.shared import Inches

methods = ["fgsm", "pgd", "pgd-linf", "pgd-linf-targ", "pgd-l2"]
dfs = ['fgsm_results_df.csv', 'pgd_results_df.csv', 'pgd_linf_results_df.csv']

for i in range(len(methods)):
    method = methods[i]
    result_df = pd.read_csv(f"results/{dfs[i]}")
    df = fgsm_results_df.copy()
    df = df.drop(['Output Image'], axis=1).reset_index(drop=True)

    # create a new Word document
    doc = docx.Document()

    # add a table to the document with the same number of rows and columns as the dataframe
    table = doc.add_table(rows=len(df)+1, cols=len(df.columns), style='Table Grid')

    # add column headers to the table
    for i, col_name in enumerate(df.columns):
        table.cell(0, i).text = col_name

    # add data and images to the table
    for i, row in df.iterrows():
        for j, col in enumerate(df.columns):
            if col == 'Output_Screenshot':
                cell = table.cell(i+1, j)
                cell._element.clear_content()
                paragraph = cell.add_paragraph()
                run = paragraph.add_run()
                run.add_picture(row[col], width=Inches(2), height=Inches(2))
            else:
                table.cell(i+1, j).text = str(row[col])

    # save the document
    doc.save(f'tables/{method}.docx')