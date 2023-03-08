# Docx Table Creator
import docx
import pandas as pd
from docx.shared import Inches

dfs = ['fgsm_results_df.csv', 'pgd_results_df.csv', 'pgd_linf_results_df.csv', 'pgd_l2_results_df.csv']

for i in range(len(dfs)):
    method = methods[i]
    df = pd.read_csv(f"results/{dfs[i]}").reset_index(drop=True)
    method = df["Method"][0]

    # create a new Word document
    doc = docx.Document()

    # add a table to the document with the same number of rows and columns as the dataframe
    table = doc.add_table(rows=len(df)+1, cols=len(df.columns), style='Table Grid')

    # add column headers to the table
    for i, col_name in enumerate(df.columns):
        table.cell(0, i).text = col_name

    # add data and images to the table
    for i, row in df.iterrows():
        for j, col in enumerate(df.columns):
            if col == 'Output_Screenshot':
                cell = table.cell(i+1, j)
                cell._element.clear_content()
                paragraph = cell.add_paragraph()
                run = paragraph.add_run()
                run.add_picture(f'{row[col]}', width=Inches(3.5), height=Inches(2))
            else:
                table.cell(i+1, j).text = str(row[col])

    # save the document
    doc.save(f'tables/{method}.docx')
    print(f'tables/{method}.docx SAVED')